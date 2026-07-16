"""
DOCX Document parser module.
Extracts raw text from Word documents using python-docx.
"""

import logging
import docx

logger = logging.getLogger(__name__)

def extract_text_from_docx(docx_path: str) -> str:
    """
    Extracts and returns plain text from a DOCX document.
    Reads both document paragraphs and structured tables to preserve block elements.
    """
    text_content = []
    
    try:
        logger.info(f"Opening DOCX file: {docx_path}")
        doc = docx.Document(docx_path)
        
        # Extract paragraph text
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text_content.append(paragraph.text)
                
        # Extract text from tables (often used for layout or skills grid in resumes)
        for table in doc.tables:
            for row in table.rows:
                row_cells_text = []
                for cell in row.cells:
                    # Strip to verify if cell has content
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in row_cells_text:
                        row_cells_text.append(cell_text)
                if row_cells_text:
                    # Join cell text in a row with a separator
                    text_content.append(" | ".join(row_cells_text))
                    
        extracted_text = "\n".join(text_content).strip()
        logger.info(f"Successfully extracted {len(extracted_text)} characters from DOCX.")
        return extracted_text
        
    except Exception as e:
        logger.error(f"DOCX extraction failed for {docx_path}: {e}")
        raise RuntimeError(f"Failed to extract text from DOCX: {e}") from e
